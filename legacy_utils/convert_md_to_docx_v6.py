import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配置
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(CURRENT_DIR, 'Project_Paper_Extended.md')
DOCX_PATH = os.path.join(CURRENT_DIR, 'Project_Paper_Final_Submission_V3.docx')
ASSETS_DIR = os.path.join(CURRENT_DIR, 'paper_assets')

# 字体大小映射
FONT_SIZES = {
    '四号': 14,
    '小四': 12,
    '五号': 10.5,
    '小五': 9
}

def set_font(run, font_name_en='Times New Roman', font_name_cn='宋体', size=12, bold=False, italic=False):
    """设置中英文字体"""
    try:
        run.font.name = font_name_en
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        r = run._element
        rPr = r.get_or_add_rPr()
        fonts = qn('w:eastAsia')
        rPr.set(fonts, font_name_cn)
    except Exception as e:
        print(f"Warning: Font setting failed: {e}")

def set_page_layout(doc):
    """设置页面布局: A4, 边距: 上下25mm, 左右20mm"""
    section = doc.sections[0]
    section.page_height = Inches(11.69) # A4 Height
    section.page_width = Inches(8.27)   # A4 Width
    section.top_margin = Inches(0.984)  # 25mm
    section.bottom_margin = Inches(0.984) # 25mm
    section.left_margin = Inches(0.787)   # 20mm
    section.right_margin = Inches(0.787)  # 20mm
    
    # 页码
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instructText = OxmlElement('w:instrText')
    instructText.set(qn('xml:space'), 'preserve')
    instructText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instructText)
    run._r.append(fldChar2)
    set_font(run, size=FONT_SIZES['小五'])

def clean_text_content(text):
    """清理Latex残余，转换为Unicode"""
    replacements = {
        r'\theta': 'θ', r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
        r'\delta': 'δ', r'\lambda': 'λ', r'\Delta': 'Δ', r'\sigma': 'σ',
        r'\sum': '∑', r'\times': '×', r'\cdot': '·', r'\sqrt': '√',
        r'\pi': 'π', r'\le': '≤', r'\ge': '≥', r'\approx': '≈',
        r'\rightarrow': '→', r'||': '‖', r'\frac{1}{2}': '1/2',
        r'_avg': 'avg', r'_{final}': 'final', 
        r'\ln': 'ln ', r'\sin': 'sin ', r'\cos': 'cos ',
        '$': '', 
        r'\frac{2\pi t}{365}': '2πt/365',
        '{': '', '}': '',
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    
    text = re.sub(r'\\frac(.*)(.*)', r'(\1)/(\2)', text)
    
    # Remove XML incompatible control characters
    # Exclude: \x09 (tab), \x0A (LF), \x0D (CR)
    # Range 0x00-0x08, 0x0B-0x0C, 0x0E-0x1F
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    
    return text

def add_paragraph_mixed(doc, text, style_settings):
    """解析混合了文本和图片的段落"""
    p = doc.add_paragraph()
    
    if 'alignment' in style_settings:
        p.alignment = style_settings['alignment']
    if 'indent' in style_settings:
        p.paragraph_format.first_line_indent = style_settings['indent']
    if 'spacing' in style_settings:
        p.paragraph_format.line_spacing = style_settings['spacing']

    pattern = r'(!\[(.*?)\]\((.*?)\))'
    parts = re.split(pattern, text)
    
    i = 0
    while i < len(parts):
        chunk = parts[i]
        
        if i + 3 < len(parts) and parts[i+1].startswith('!['):
            if chunk:
                add_text_run(p, chunk, style_settings)
            
            img_alt = parts[i+2]
            img_path_raw = parts[i+3]
            
            img_path = os.path.join(CURRENT_DIR, img_path_raw.replace('/', os.sep).replace('\\', os.sep))
            if not os.path.exists(img_path):
                img_path = os.path.join(ASSETS_DIR, os.path.basename(img_path_raw))
            
            if os.path.exists(img_path):
                run = p.add_run()
                is_block_image = (len(text.strip()) == len(parts[i+1].strip())) or (text.strip() == parts[i+1])
                
                try:
                    if is_block_image:
                        if 'eq_' in os.path.basename(img_path):
                             run.add_picture(img_path, width=Inches(3.5))
                        else:
                             run.add_picture(img_path, width=Inches(5.0))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.first_line_indent = Inches(0)
                    else:
                        run.add_picture(img_path, height=Pt(14)) 
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
            else:
                run = p.add_run(f"[缺失图片: {img_alt}]")
                run.font.color.rgb = RGBColor(255, 0, 0)

            i += 4
        else:
            if chunk:
                add_text_run(p, chunk, style_settings)
            i += 1

def add_text_run(p, text, style_settings):
    """添加纯文本 run, 处理粗体和清理"""
    text = clean_text_content(text)
    
    bold_parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in bold_parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = p.add_run(content)
            set_font(run, 'Times New Roman', style_settings['cn_font'], style_settings['size'], bold=True)
        else:
            if part:
                run = p.add_run(part)
                set_font(run, 'Times New Roman', style_settings['cn_font'], style_settings['size'], bold=style_settings['bold'])

def create_docx(lines, output_path):
    doc = Document()
    set_page_layout(doc)
    
    # 样式配置
    H1_Style = {'size': FONT_SIZES['四号'], 'cn_font': '黑体', 'bold': True, 'spacing': 1.0}
    H2_Style = {'size': FONT_SIZES['小四'], 'cn_font': '黑体', 'bold': True, 'spacing': 1.0}
    Body_Style = {'size': FONT_SIZES['小四'], 'cn_font': '宋体', 'bold': False, 'spacing': 1.0, 'indent': Inches(0.3)}
    Caption_Style = {'size': FONT_SIZES['五号'], 'cn_font': '宋体', 'bold': False, 'alignment': WD_ALIGN_PARAGRAPH.CENTER}
    Title_Style = {'size': 18, 'cn_font': '黑体', 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER}

    for line in lines:
        line = line.strip()
        if not line: continue
        
        if line.startswith('---') or (len(line) > 3 and line == '-'*len(line)):
            continue

        # 1. 标题
        if line.startswith('# '):
            add_paragraph_mixed(doc, line[2:], Title_Style)
        elif line.startswith('## '):
            add_paragraph_mixed(doc, line[3:], H1_Style)
        elif line.startswith('### '):
            add_paragraph_mixed(doc, line[4:], H2_Style)
            
        # 2. 题注 *图...*
        elif line.startswith('*图') and line.endswith('*'):
            add_paragraph_mixed(doc, line.strip('*'), Caption_Style)
        
        # 3. 列表 (自动识别 1. 开头)
        elif re.match(r'^\d+\.', line):
            List_Style = Body_Style.copy()
            List_Style['indent'] = Inches(0)
            add_paragraph_mixed(doc, line, List_Style)
            
        # 4. 普通符号列表 - 强制去圆点
        elif line.startswith('* ') or line.startswith('- '):
             # 替换为数字? 或者只用文本
             # 既然MD里已经全改了，这里应该是漏网之鱼
             # 用 '- ' 开头，不加圆点
             clean_line = line[2:] 
             List_Style = Body_Style.copy()
             List_Style['indent'] = Inches(0)
             # add generic prefix to indicate list without dot
             # clean_line = "  " + clean_line
             add_paragraph_mixed(doc, clean_line, List_Style)

        # 5. 普通段落
        else:
            add_paragraph_mixed(doc, line, Body_Style)

    try:
        doc.save(output_path)
        print(f"DONE: {output_path}")
    except PermissionError:
        print(f"ERROR: Permission denied. Please close code/Word file: {output_path}")

if __name__ == "__main__":
    if os.path.exists(MD_PATH):
        with open(MD_PATH, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        create_docx(lines, DOCX_PATH)
    else:
        print("MD file not found")
